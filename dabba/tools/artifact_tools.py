"""
Document generation tools for the dabba agent.

file_write only ever produces plain text — it has no PDF/DOCX rendering
logic, so a model calling it with a ".pdf" path silently writes a text
file with a misleading extension and no way to detect the mismatch.
These tools render real binary documents from markdown so "create a PDF"
requests produce an actual PDF instead of a renamed text file.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

from dabba.agent.tool_schema import ToolDefinition, ToolParameter
from dabba.agent.tool_registry import ToolRegistry
from dabba.utils.logging import get_logger

logger = get_logger("dabba.tools.artifact_tools")


@dataclass
class _Block:
    kind: str  # "heading" | "code" | "list" | "paragraph"
    text: str
    level: int = 0
    ordered: bool = False


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)")


def _parse_markdown(markdown_text: str) -> List[_Block]:
    """
    Parse a practical subset of markdown into renderable blocks: headings,
    fenced code blocks, bullet/numbered lists, and paragraphs.

    Not a full CommonMark parser — inline emphasis/links pass through as
    literal text. That's enough for the kind of generated docs (summaries,
    READMEs, reports) these tools exist to render.
    """
    blocks: List[_Block] = []
    lines = markdown_text.splitlines()
    i = 0
    paragraph_buf: List[str] = []

    def flush_paragraph() -> None:
        if paragraph_buf:
            blocks.append(_Block(kind="paragraph", text=" ".join(paragraph_buf)))
            paragraph_buf.clear()

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            flush_paragraph()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(_Block(kind="code", text="\n".join(code_lines)))
            i += 1  # skip closing fence
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            flush_paragraph()
            blocks.append(
                _Block(kind="heading", text=heading.group(2).strip(), level=len(heading.group(1)))
            )
            i += 1
            continue

        bullet = _BULLET_RE.match(line)
        numbered = _NUMBERED_RE.match(line)
        if bullet or numbered:
            flush_paragraph()
            text = (bullet or numbered).group(1).strip()
            blocks.append(_Block(kind="list", text=text, ordered=bool(numbered)))
            i += 1
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        paragraph_buf.append(line.strip())
        i += 1

    flush_paragraph()
    return blocks


def markdown_to_pdf(path: str, markdown_text: str, title: str = "") -> Dict[str, object]:
    """
    Render markdown text to a real PDF file using reportlab.

    Args:
        path: Output path for the .pdf file (extension corrected if missing/wrong).
        markdown_text: Markdown-formatted document content.
        title: Optional document title, rendered as the first heading.

    Returns:
        Dict with status/path/size, verified against the filesystem after
        the render — not just assumed because doc.build() didn't raise.

    Raises:
        OSError: If the PDF cannot be written, or the render produced no file.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
    )

    file_path = Path(path).expanduser().resolve()
    if file_path.suffix.lower() != ".pdf":
        file_path = file_path.with_suffix(".pdf")
    file_path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(file_path), pagesize=LETTER, title=title or file_path.stem)
    story = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 0.2 * inch))

    heading_styles = {1: "Heading1", 2: "Heading2", 3: "Heading3"}
    list_buf: List[ListItem] = []
    list_ordered = False

    def flush_list() -> None:
        nonlocal list_buf, list_ordered
        if list_buf:
            story.append(ListFlowable(list_buf, bulletType="1" if list_ordered else "bullet"))
            story.append(Spacer(1, 0.1 * inch))
            list_buf = []

    for block in _parse_markdown(markdown_text):
        if block.kind == "list":
            list_ordered = block.ordered
            list_buf.append(ListItem(Paragraph(block.text, styles["Normal"])))
            continue
        flush_list()
        if block.kind == "heading":
            story.append(Paragraph(block.text, styles[heading_styles.get(block.level, "Heading4")]))
        elif block.kind == "code":
            story.append(Preformatted(block.text, styles["Code"]))
        else:
            story.append(Paragraph(block.text, styles["Normal"]))
        story.append(Spacer(1, 0.12 * inch))
    flush_list()

    if not story:
        story.append(Paragraph("(empty document)", styles["Normal"]))

    doc.build(story)

    if not file_path.exists() or file_path.stat().st_size == 0:
        raise OSError(f"PDF render reported no error but produced no file: {file_path}")

    size = file_path.stat().st_size
    logger.info("Wrote PDF '%s' (%d bytes)", file_path, size)
    return {"status": "success", "path": str(file_path), "size": size, "exists": True}


def markdown_to_docx(path: str, markdown_text: str, title: str = "") -> Dict[str, object]:
    """
    Render markdown text to a real Word (.docx) file using python-docx.

    Args:
        path: Output path for the .docx file (extension corrected if missing/wrong).
        markdown_text: Markdown-formatted document content.
        title: Optional document title, rendered as the first heading.

    Returns:
        Dict with status/path/size, verified against the filesystem after
        the save — not just assumed because document.save() didn't raise.

    Raises:
        OSError: If the DOCX cannot be written, or the render produced no file.
    """
    import docx as docx_lib

    file_path = Path(path).expanduser().resolve()
    if file_path.suffix.lower() != ".docx":
        file_path = file_path.with_suffix(".docx")
    file_path.parent.mkdir(parents=True, exist_ok=True)

    document = docx_lib.Document()
    if title:
        document.add_heading(title, level=0)

    heading_levels = {1: 1, 2: 2, 3: 3, 4: 4, 5: 5, 6: 6}
    for block in _parse_markdown(markdown_text):
        if block.kind == "heading":
            document.add_heading(block.text, level=heading_levels.get(block.level, 6))
        elif block.kind == "code":
            paragraph = document.add_paragraph(block.text)
            for run in paragraph.runs:
                run.font.name = "Courier New"
        elif block.kind == "list":
            document.add_paragraph(block.text, style="List Number" if block.ordered else "List Bullet")
        else:
            document.add_paragraph(block.text)

    document.save(str(file_path))

    if not file_path.exists() or file_path.stat().st_size == 0:
        raise OSError(f"DOCX render reported no error but produced no file: {file_path}")

    size = file_path.stat().st_size
    logger.info("Wrote DOCX '%s' (%d bytes)", file_path, size)
    return {"status": "success", "path": str(file_path), "size": size, "exists": True}


def register_artifact_tools(registry: ToolRegistry) -> None:
    """
    Register markdown-to-PDF/DOCX document generation tools with a ToolRegistry.

    Args:
        registry: The tool registry to register with.
    """
    registry.register(
        ToolDefinition(
            name="markdown_to_pdf",
            description=(
                "Render markdown text to a real PDF file on disk. Use this instead of "
                "file_write for '.pdf' paths — file_write only writes plain text and "
                "cannot produce an actual PDF."
            ),
            parameters=[
                ToolParameter(name="path", type="string", description="Output path for the .pdf file."),
                ToolParameter(
                    name="markdown_text", type="string", description="Markdown-formatted document content."
                ),
                ToolParameter(
                    name="title", type="string", description="Optional document title.",
                    required=False, default="",
                ),
            ],
            handler=markdown_to_pdf,
            handler_sync=True,
            category="filesystem",
        )
    )
    registry.register(
        ToolDefinition(
            name="markdown_to_docx",
            description=(
                "Render markdown text to a real Word (.docx) file on disk. Use this instead "
                "of file_write for '.docx' paths — file_write only writes plain text and "
                "cannot produce an actual Word document."
            ),
            parameters=[
                ToolParameter(name="path", type="string", description="Output path for the .docx file."),
                ToolParameter(
                    name="markdown_text", type="string", description="Markdown-formatted document content."
                ),
                ToolParameter(
                    name="title", type="string", description="Optional document title.",
                    required=False, default="",
                ),
            ],
            handler=markdown_to_docx,
            handler_sync=True,
            category="filesystem",
        )
    )
